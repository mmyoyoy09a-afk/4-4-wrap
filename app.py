import os
import re
import sqlite3
from datetime import datetime, timezone
from functools import wraps
import pytz  # Handles target timezone offsets perfectly
from flask import Flask, render_template, request, redirect, url_for, session, g, abort, flash

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-me")

DB_PATH = os.environ.get("DB_PATH", "data/app.db")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "changeme")

# Hardcoded reveal date: August 1, 2026, at 00:00:00 PHT (UTC+8)
REVEAL_AT_RAW = "2026-08-01T00:00:00+08:00"

# ---------- Validation block ----------
try:
    parsed_date = datetime.fromisoformat(REVEAL_AT_RAW)
    utc_date = parsed_date.astimezone(timezone.utc)
    now_with_tz = datetime.now(timezone.utc)
    is_past = now_with_tz >= utc_date

    print("--- Timezone Parser Test ---")
    print(f"Parsed Local Time: {parsed_date}")
    print(f"Timezone Offset:   {parsed_date.tzinfo}")
    print(f"Equivalent UTC:    {utc_date}")
    print(f"Current UTC Time:  {now_with_tz}")
    print(f"Has date passed?:  {is_past} (True means messages are now unlocked)")

except ValueError as e:
    print(f"Parsing Error: Your date string is invalid. Details: {e}")

# ---------- Reveal Utility Functions ----------
def get_reveal_dt():
    if not REVEAL_AT_RAW:
        return None
    try:
        dt = datetime.fromisoformat(REVEAL_AT_RAW)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt
    except ValueError:
        return None

def is_revealed():
    """Checks if the current UTC time has passed the target reveal UTC time."""
    target_dt = get_reveal_dt()
    if not target_dt:
        return False
    
    # Compare both values securely in UTC space
    now_utc = datetime.now(timezone.utc)
    target_utc = target_dt.astimezone(timezone.utc)
    return now_utc >= target_utc

def get_reveal_display():
    """Returns status text showing when messages unlock or if they are open."""
    target_dt = get_reveal_dt()
    if not target_dt:
        return "No reveal date set."
        
    if is_revealed():
        return "Messages have been officially unlocked!"
        
    return f"Messages will unlock on: {target_dt.strftime('%B %d, %Y %I:%M %p %Z')}"



# ---------- card export (.docx) helpers ----------

FONT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts", "Baloo2-Regular.ttf")
_font_cache = {}


def _get_font(weight, size):
    from PIL import ImageFont
    key = (weight, size)
    if key in _font_cache:
        return _font_cache[key]
    font = ImageFont.truetype(FONT_PATH, size)
    try:
        font.set_variation_by_axes([weight])
    except Exception:
        pass
    _font_cache[key] = font
    return font


def _hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def _darken(hex_color, factor=0.8):
    r, g, b = _hex_to_rgb(hex_color)
    return (int(r * factor), int(g * factor), int(b * factor))


def _wrap_text(draw, text, font, max_width):
    lines = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        cur = ""
        for w in words:
            trial = (cur + " " + w).strip()
            if draw.textlength(trial, font=font) <= max_width or not cur:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        lines.append(cur)
    return lines


def _draw_capsule(draw, p1, p2, width, color):
    draw.line([p1, p2], fill=color, width=width)
    r = width / 2
    for p in (p1, p2):
        draw.ellipse([p[0] - r, p[1] - r, p[0] + r, p[1] + r], fill=color)


def _rounded_rect_points(x0, y0, x1, y1, r, seg_len=9):
    import math

    def arc(cx, cy, rad, a0, a1):
        n = max(2, int(abs(a1 - a0) / 8))
        return [
            (cx + rad * math.cos(math.radians(a0 + (a1 - a0) * i / n)),
             cy + rad * math.sin(math.radians(a0 + (a1 - a0) * i / n)))
            for i in range(n + 1)
        ]

    def subdiv(p0, p1):
        d = math.hypot(p1[0] - p0[0], p1[1] - p0[1])
        n = max(1, int(d / seg_len))
        return [(p0[0] + (p1[0] - p0[0]) * i / n, p0[1] + (p1[1] - p0[1]) * i / n) for i in range(n)]

    pts = []
    pts += subdiv((x0 + r, y0), (x1 - r, y0))
    pts += arc(x1 - r, y0 + r, r, -90, 0)
    pts += subdiv((x1, y0 + r), (x1, y1 - r))
    pts += arc(x1 - r, y1 - r, r, 0, 90)
    pts += subdiv((x1 - r, y1), (x0 + r, y1))
    pts += arc(x0 + r, y1 - r, r, 90, 180)
    pts += subdiv((x0, y1 - r), (x0, y0 + r))
    pts += arc(x0 + r, y0 + r, r, 180, 270)
    pts.append(pts[0])
    return pts


def _draw_dashed_rounded_rect(draw, x0, y0, x1, y1, r, color, width=7, dash_segments=3, gap_segments=2):
    pts = _rounded_rect_points(x0, y0, x1, y1, r)
    n = len(pts)
    i = 0
    while i < n - 1:
        end = min(i + dash_segments, n - 1)
        _draw_capsule(draw, pts[i], pts[end], width, color)
        i = end + gap_segments


def render_card_image(text, to_name, from_name, color_hex, page, total, w=1080, h=720):
    from PIL import Image, ImageDraw
    import io

    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    fill = _hex_to_rgb(color_hex)
    border = _darken(color_hex, 0.8)
    radius = 46
    pad = 14
    draw.rounded_rectangle([pad, pad, w - pad, h - pad], radius=radius, fill=fill)
    _draw_dashed_rounded_rect(draw, pad, pad, w - pad, h - pad, radius, border)

    text_color = (44, 74, 82)
    muted_color = (124, 154, 158)
    pad_in = 46

    f_small = _get_font(600, 20)
    draw.text((pad_in, pad_in - 8), f"({page}/{total})", font=f_small, fill=muted_color)

    y = pad_in + 30
    if to_name:
        f_to = _get_font(700, 55)
        draw.text((pad_in, y), f"To: {to_name}", font=f_to, fill=text_color)
        y += 75
    else:
        y += 6

    from_line_h = 75 if from_name else 0
    f_msg = _get_font(500, 55)
    max_w = w - 2 * pad_in
    lines = _wrap_text(draw, text, f_msg, max_w)
    line_h = 90
    max_lines = max(1, int((h - y - pad_in - from_line_h) / line_h))
    for line in lines[:max_lines]:
        draw.text((pad_in, y), line, font=f_msg, fill=text_color)
        y += line_h

    if from_name:
        f_from = _get_font(600, 55)
        label = f"From: {from_name}"
        tw = draw.textlength(label, font=f_from)
        draw.text((w - pad_in - tw, h - pad_in - 48), label, font=f_from, fill=text_color)

    bio = io.BytesIO()
    img.save(bio, format="PNG")
    bio.seek(0)
    return bio


def _set_table_borders_none(table):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_margins_zero(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _set_cell_width(cell, cm_width):
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cell.width = Cm(cm_width)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm_width * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def build_cards_docx(rows, pages_by_card):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.section import WD_ORIENT
    import io

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    narrow = Cm(1.27)
    section.top_margin = narrow
    section.bottom_margin = narrow
    section.left_margin = narrow
    section.right_margin = narrow

    # flatten each card into N grid "slots", one per page
    slots = []
    for r in rows:
        pages = pages_by_card.get(r["id"]) or []
        total = len(pages)
        for i, text in enumerate(pages):
            page_num = i + 1
            is_first = page_num == 1
            is_last = page_num == total
            slots.append({
                "color": r["color"], "text": text, "page": page_num, "total": total,
                "to": r["to_name"] if is_first else None,
                "from": r["from_name"] if is_last else None,
            })

    for start in range(0, len(slots), 8):
        chunk = slots[start:start + 8]
        table = doc.add_table(rows=4, cols=2)
        table.autofit = False
        _set_table_borders_none(table)
        for r_i in range(4):
            table.rows[r_i].height = Cm(CARD_H_CM)
            table.rows[r_i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        for idx, slot in enumerate(chunk):
            r_i, c_i = divmod(idx, 2)
            cell = table.cell(r_i, c_i)
            _set_cell_width(cell, CARD_W_CM)
            _set_cell_margins_zero(cell)
            cell.text = ""

            png_bio = render_card_image(
                slot["text"], slot["to"], slot["from"], slot["color"], slot["page"], slot["total"]
            )
            run = cell.paragraphs[0].add_run()
            run.add_picture(png_bio, width=Cm(CARD_W_CM), height=Cm(CARD_H_CM))

        if start + 8 < len(slots):
            doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

SCHEMA = """
CREATE TABLE IF NOT EXISTS people (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL,
    slug TEXT NOT NULL UNIQUE,
    created_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS messages (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    person_id INTEGER NOT NULL,
    sender_name TEXT,
    content TEXT NOT NULL,
    created_at TEXT NOT NULL,
    FOREIGN KEY (person_id) REFERENCES people (id)
);

CREATE TABLE IF NOT EXISTS cards (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    to_name TEXT,
    from_name TEXT,
    color TEXT NOT NULL,
    page1_text TEXT,
    page2_text TEXT,
    show_now INTEGER NOT NULL DEFAULT 0,
    created_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS card_pages (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    card_id INTEGER NOT NULL,
    page_number INTEGER NOT NULL,
    text TEXT NOT NULL,
    FOREIGN KEY (card_id) REFERENCES cards (id)
);

CREATE TABLE IF NOT EXISTS settings (
    key TEXT PRIMARY KEY,
    value TEXT NOT NULL
);
"""

CARD_COLORS = [
    {"hex": "#f7e5ec", "name": "Pink", "meaning": "Love and Deep Platonic Bonds"},
    {"hex": "#fcf6bd", "name": "Pastel Yellow", "meaning": "Admiration and Recognition"},
    {"hex": "#d7ffce", "name": "Mint Green", "meaning": "Grace and Gentle Reflection"},
    {"hex": "#e1faff", "name": "Soft Blue", "meaning": "Prosper and Godspeed"},
    {"hex": "#e9d5e7", "name": "Pale Lavender", "meaning": "Gratitude and Appreciation"},
    {"hex": "#f8f7ff", "name": "Pure White", "meaning": "Leave it Blank"}
]
CARD_COLOR_NAMES = {c["hex"]: c["name"] for c in CARD_COLORS}

CARD_W_CM = 9  # landscape: fits a strict 2-col x 4-row, 8-per-page grid on A4
CARD_H_CM = 6


def get_db():
    if "db" not in g:
        os.makedirs(os.path.dirname(DB_PATH) or ".", exist_ok=True)
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
        g.db.execute("PRAGMA foreign_keys = ON")
    return g.db


@app.teardown_appcontext
def close_db(exception=None):
    db = g.pop("db", None)
    if db is not None:
        db.close()


def init_db():
    os.makedirs(os.path.dirname(DB_PATH) or ".", exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.executescript(SCHEMA)
    conn.commit()

    # safe migration: add show_now to older deployed DBs that predate it
    cols = [r[1] for r in conn.execute("PRAGMA table_info(cards)").fetchall()]
    if "show_now" not in cols:
        conn.execute("ALTER TABLE cards ADD COLUMN show_now INTEGER NOT NULL DEFAULT 0")
        conn.commit()

    # migrate any old page1_text/page2_text rows into card_pages (once)
    rows = conn.execute("SELECT id, page1_text, page2_text FROM cards").fetchall()
    for card_id, p1, p2 in rows:
        already = conn.execute(
            "SELECT COUNT(*) FROM card_pages WHERE card_id = ?", (card_id,)
        ).fetchone()[0]
        if already == 0 and p1:
            conn.execute(
                "INSERT INTO card_pages (card_id, page_number, text) VALUES (?, 1, ?)", (card_id, p1)
            )
            if p2:
                conn.execute(
                    "INSERT INTO card_pages (card_id, page_number, text) VALUES (?, 2, ?)", (card_id, p2)
                )
    conn.commit()
    conn.close()


def get_card_pages(db, card_id):
    rows = db.execute(
        "SELECT text FROM card_pages WHERE card_id = ? ORDER BY page_number ASC", (card_id,)
    ).fetchall()
    return [r["text"] for r in rows]


def get_setting(db, key, default=None):
    row = db.execute("SELECT value FROM settings WHERE key = ?", (key,)).fetchone()
    return row["value"] if row else default


def set_setting(db, key, value):
    db.execute("DELETE FROM settings WHERE key = ?", (key,))
    db.execute("INSERT INTO settings (key, value) VALUES (?, ?)", (key, value))
    db.commit()


def allow_hide_choice(db):
    return get_setting(db, "allow_hide_choice", "1") == "1"


def slugify(name):
    slug = name.strip().lower()
    slug = re.sub(r"[^a-z0-9]+", "-", slug)
    slug = slug.strip("-")
    return slug or "member"


def unique_slug(db, name):
    base = slugify(name)
    slug = base
    n = 2
    while db.execute("SELECT 1 FROM people WHERE slug = ?", (slug,)).fetchone():
        slug = f"{base}-{n}"
        n += 1
    return slug


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not session.get("is_admin"):
            return redirect(url_for("admin_login", next=request.path))
        return view(*args, **kwargs)
    return wrapped


# ---------- public routes ----------

@app.route("/")
def index():
    db = get_db()
    raw_cards = db.execute("SELECT * FROM cards ORDER BY id DESC").fetchall()
    global_revealed = is_revealed()
    cards = []
    for c in raw_cards:
        row = dict(c)
        card_revealed = global_revealed or bool(row.get("show_now"))
        pages = get_card_pages(db, row["id"])
        row["total_pages"] = len(pages)
        row["revealed"] = card_revealed
        if card_revealed:
            row["pages"] = pages
            row["preview_text"] = pages[0] if pages else ""
        else:
            row["pages"] = []
            row["preview_text"] = None
        cards.append(row)
    any_hidden = any(not c["revealed"] for c in cards)
    return render_template(
        "index.html", cards=cards, revealed=global_revealed, any_hidden=any_hidden,
        reveal_display=get_reveal_display()
    )


@app.route("/p/<slug>", methods=["GET"])
def person(slug):
    db = get_db()
    person = db.execute("SELECT * FROM people WHERE slug = ?", (slug,)).fetchone()
    if not person:
        abort(404)
    raw_messages = db.execute(
        "SELECT * FROM messages WHERE person_id = ? ORDER BY id DESC", (person["id"],)
    ).fetchall()

    revealed = is_revealed()
    messages = []
    for m in raw_messages:
        row = dict(m)
        if not revealed:
            row["content"] = None  # withhold real text until reveal date
        messages.append(row)

    return render_template(
        "person.html",
        person=person,
        messages=messages,
        revealed=revealed,
        reveal_display=get_reveal_display(),
    )


@app.route("/p/<slug>/message", methods=["POST"])
def send_message(slug):
    db = get_db()
    person = db.execute("SELECT * FROM people WHERE slug = ?", (slug,)).fetchone()
    if not person:
        abort(404)

    content = (request.form.get("content") or "").strip()
    is_anon = request.form.get("is_anon") == "on"
    sender_name = (request.form.get("sender_name") or "").strip()

    if not content:
        flash("Message can't be empty.", "error")
        return redirect(url_for("person", slug=slug))

    if is_anon or not sender_name:
        sender_name = None

    db.execute(
        "INSERT INTO messages (person_id, sender_name, content, created_at) VALUES (?, ?, ?, ?)",
        (person["id"], sender_name, content, datetime.utcnow().isoformat()),
    )
    db.commit()
    flash("Message sent.", "success")
    return redirect(url_for("person", slug=slug))


# ---------- card designer (public) ----------

@app.route("/cards/new", methods=["GET"])
def card_new():
    db = get_db()
    return render_template("card_new.html", colors=CARD_COLORS, allow_hide=allow_hide_choice(db))


@app.route("/cards/new", methods=["POST"])
def card_create():
    data = request.get_json(silent=True) or {}
    to_name = (data.get("to_name") or "").strip()[:100]
    from_name = (data.get("from_name") or "").strip()[:100]
    color = (data.get("color") or "").strip()
    raw_pages = data.get("pages")
    show_now = bool(data.get("show_now"))

    if not isinstance(raw_pages, list):
        raw_pages = []
    pages = [p.strip() for p in raw_pages if isinstance(p, str) and p.strip()]

    if color not in CARD_COLOR_NAMES:
        return {"ok": False, "error": "Invalid color."}, 400
    if not pages:
        return {"ok": False, "error": "Message can't be empty."}, 400

    db = get_db()
    if not allow_hide_choice(db):
        show_now = True  # hiding is disabled site-wide; every card posts visible immediately

    cur = db.execute(
        """INSERT INTO cards (to_name, from_name, color, show_now, created_at)
           VALUES (?, ?, ?, ?, ?)""",
        (to_name or None, from_name or None, color, 1 if show_now else 0, datetime.utcnow().isoformat()),
    )
    card_id = cur.lastrowid
    for i, text in enumerate(pages, start=1):
        db.execute(
            "INSERT INTO card_pages (card_id, page_number, text) VALUES (?, ?, ?)", (card_id, i, text)
        )
    db.commit()
    return {"ok": True}


# ---------- admin routes ----------

@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        password = request.form.get("password", "")
        if password == ADMIN_PASSWORD:
            session["is_admin"] = True
            next_url = request.form.get("next") or url_for("admin_cards")
            return redirect(next_url)
        flash("Wrong password.", "error")
    return render_template("admin_login.html", next=request.args.get("next", ""))


@app.route("/admin/logout")
def admin_logout():
    session.pop("is_admin", None)
    return redirect(url_for("index"))


@app.route("/admin")
@login_required
def admin():
    db = get_db()
    people = db.execute(
        """SELECT p.*, COUNT(m.id) as msg_count
           FROM people p LEFT JOIN messages m ON m.person_id = p.id
           GROUP BY p.id ORDER BY p.name COLLATE NOCASE ASC"""
    ).fetchall()
    recent_messages = db.execute(
        """SELECT m.*, p.name as person_name, p.slug as person_slug
           FROM messages m JOIN people p ON p.id = m.person_id
           ORDER BY m.id DESC LIMIT 50"""
    ).fetchall()
    return render_template(
        "admin.html",
        people=people,
        recent_messages=recent_messages,
        revealed=is_revealed(),
        reveal_display=get_reveal_display(),
    )


@app.route("/admin/people/add", methods=["POST"])
@login_required
def add_person():
    db = get_db()
    name = (request.form.get("name") or "").strip()
    if not name:
        flash("Name can't be empty.", "error")
        return redirect(url_for("admin"))
    slug = unique_slug(db, name)
    db.execute(
        "INSERT INTO people (name, slug, created_at) VALUES (?, ?, ?)",
        (name, slug, datetime.utcnow().isoformat()),
    )
    db.commit()
    flash(f"Added {name}.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/people/<int:person_id>/delete", methods=["POST"])
@login_required
def delete_person(person_id):
    db = get_db()
    db.execute("DELETE FROM messages WHERE person_id = ?", (person_id,))
    db.execute("DELETE FROM people WHERE id = ?", (person_id,))
    db.commit()
    flash("Removed.", "success")
    return redirect(url_for("admin"))


@app.route("/admin/messages/<int:message_id>/delete", methods=["POST"])
@login_required
def delete_message(message_id):
    db = get_db()
    db.execute("DELETE FROM messages WHERE id = ?", (message_id,))
    db.commit()
    flash("Message removed.", "success")
    return redirect(url_for("admin"))


# ---------- admin: card management ----------

@app.route("/admin/cards")
@login_required
def admin_cards():
    db = get_db()
    raw_cards = db.execute("SELECT * FROM cards ORDER BY id DESC").fetchall()
    cards = []
    for c in raw_cards:
        row = dict(c)
        row["pages"] = get_card_pages(db, row["id"])
        cards.append(row)
    return render_template(
        "admin_cards.html", cards=cards, color_names=CARD_COLOR_NAMES, allow_hide=allow_hide_choice(db)
    )


@app.route("/admin/settings", methods=["POST"])
@login_required
def update_settings():
    db = get_db()
    value = "1" if request.form.get("allow_hide_choice") == "on" else "0"
    set_setting(db, "allow_hide_choice", value)
    flash("Setting saved.", "success")
    return redirect(url_for("admin_cards"))


@app.route("/admin/cards/<int:card_id>/delete", methods=["POST"])
@login_required
def delete_card(card_id):
    db = get_db()
    db.execute("DELETE FROM card_pages WHERE card_id = ?", (card_id,))
    db.execute("DELETE FROM cards WHERE id = ?", (card_id,))
    db.commit()
    flash("Card removed.", "success")
    return redirect(url_for("admin_cards"))


@app.route("/admin/cards/export", methods=["POST"])
@login_required
def export_cards():
    from flask import send_file

    ids = request.form.getlist("card_ids")
    if not ids:
        flash("Select at least one card to export.", "error")
        return redirect(url_for("admin_cards"))

    db = get_db()
    placeholders = ",".join("?" for _ in ids)
    rows = db.execute(
        f"SELECT * FROM cards WHERE id IN ({placeholders}) ORDER BY id ASC", ids
    ).fetchall()
    pages_by_card = {r["id"]: get_card_pages(db, r["id"]) for r in rows}

    bio = build_cards_docx(rows, pages_by_card)
    return send_file(
        bio,
        as_attachment=True,
        download_name="bsece44-cards.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


init_db()

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
